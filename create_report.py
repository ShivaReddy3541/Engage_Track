import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_specification_document():
    doc = Document()

    # 1. Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styling helper variables
    COLOR_PRIMARY = RGBColor(26, 54, 93)   # Deep Navy
    COLOR_SECONDARY = RGBColor(0, 150, 166) # Teal
    COLOR_TEXT = RGBColor(51, 65, 85)       # Slate Grey
    COLOR_MUTED = RGBColor(100, 116, 139)   # Cool Grey
    
    # 2. Helpers for XML layout styling (shading & borders)
    def set_cell_bg(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
        # 1 dxa = 1/20 of a point. 140 dxa = ~7pt, 180 dxa = ~9pt padding
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{name}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_cell_borders(cell, top="D3D3D3", bottom="D3D3D3", left=None, right=None):
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        borders = [('top', top), ('bottom', bottom), ('left', left), ('right', right)]
        for name, hex_val in borders:
            if hex_val:
                node = OxmlElement(f'w:{name}')
                node.set(qn('w:val'), 'single')
                node.set(qn('w:sz'), '4')  # 1/8 pt width
                node.set(qn('w:space'), '0')
                node.set(qn('w:color'), hex_val)
                tcBorders.append(node)
            else:
                node = OxmlElement(f'w:{name}')
                node.set(qn('w:val'), 'none')
                tcBorders.append(node)
        tcPr.append(tcBorders)

    # 3. Document Styling Defaults
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # 4. Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(72)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("EngageAI: Multi-Agent Intelligent LMS Portal")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    run_sub = sub_p.add_run("System Design and Internal Architecture Specification")
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = COLOR_MUTED

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_before = Pt(144)
    run_meta = meta_p.add_run("Prepared for Project Report Portfolio\nAcademic Year 2026")
    run_meta.font.size = Pt(11)
    run_meta.font.italic = True
    run_meta.font.color.rgb = COLOR_MUTED

    doc.add_page_break()

    # 5. Helper Function to Add Styled Headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_bg(cell, "F3F4F6") # Light grey background
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        set_cell_borders(cell, top="E2E8F0", bottom="E2E8F0", left="3182CE", right="E2E8F0") # Blue accent border on left
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 41, 59)
        # Empty space after table
        space = doc.add_paragraph()
        space.paragraph_format.space_before = Pt(0)
        space.paragraph_format.space_after = Pt(6)

    # 6. Writing content sections
    # ----------------------------------------------------
    # SECTION 1
    # ----------------------------------------------------
    add_heading_1("1. Overall Architecture")
    
    doc.add_paragraph(
        "The EngageAI Multi-Agent LMS Portal is built on a highly decoupled, state-of-the-art microservice "
        "architecture consisting of seven independently deployable services: one user-facing React web frontend, "
        "one core FastAPI backend service, and five autonomous cognitive agent microservices (including the central "
        "Orchestrator). The ecosystem utilizes a centralized PostgreSQL database for relational, structured academic records "
        "and a high-performance Redis container acting as the message broker for real-time pub/sub event notifications."
    )
    
    doc.add_paragraph(
        "To ensure development consistency and prevent language-switching overhead, all backend microservices are "
        "implemented in Python. This allows machine learning pipelines, OpenCV tasks, and NLP modules to share "
        "hashing utilities, validation routers, and models seamlessly. The frontend client acts strictly as a "
        "Single-Page Application (SPA), querying core backend resources over secure HTTPS APIs or receiving real-time "
        "notifications via WebSockets, ensuring zero direct database exposures."
    )

    doc.add_paragraph(
        "Relational CRUD operations and authentication protocols utilize traditional REST HTTP protocols. In contrast, "
        "agent alerts and tracking updates are managed asynchronously via Redis Channels to prevent heavy HTTP polling cycles. "
        "The diagram below illustrates this decoupled communication model:"
    )

    add_code_block(
        "+-------------------+             REST HTTPS            +-------------------+\n"
        "|  React Frontend   | <===============================> |  Core FastAPI API  |\n"
        "|  (Vite + Tailwind)|                                   |  (Database CRUD)   |\n"
        "+-------------------+                                   +-------------------+\n"
        "         ^                                                        ||\n"
        "         || WebSockets Gateway (Real-Time Alerts)                 ||\n"
        "         v                                                        v\n"
        "+-------------------+           Asynchronous Events     +-------------------+\n"
        "|   Orchestrator    | <-------------------------------- |  Postgres + Redis |\n"
        "|   Rules Engine    |       [ Redis Channels Pub/Sub ]  |  (Ledgers & Bus)  |\n"
        "+-------------------+                                   +-------------------+\n"
        "         ^                                                        ^ \n"
        "         |========================================================|\n"
        "         |      Absence / EAR / Screen YOLO / Content Moderation Agents"
    )

    # ----------------------------------------------------
    # SECTION 2
    # ----------------------------------------------------
    add_heading_1("2. Frontend Client Architecture")
    
    doc.add_paragraph(
        "The frontend client is engineered using React 18 and Vite for optimized module bundlings. Tailwind CSS is "
        "employed to implement a dynamic, white-themed responsive split-card layout based on MeritCurve guidelines, "
        "ensuring unified interfaces across desktops, tablets, and phones."
    )

    add_heading_2("Authentication Storage and Security")
    doc.add_paragraph(
        "Upon successful user verification at the endpoint POST /auth/login, the backend returns a stateless "
        "JSON Web Token (JWT). To shield the system from Cross-Site Scripting (XSS) vulnerabilities, the token is saved "
        "exclusively in React state memory rather than browser LocalStorage. Short-lived access tokens (60 minutes) "
        "accompanied by automated Axios middleware interceptors handle refreshing routines. If a service request yields "
        "a 401 Unauthorized status, the client intercepts the call and routes the user back to the login gateway."
    )

    add_heading_2("Role-Based UI Layout Dispatcher")
    doc.add_paragraph(
        "Decoded JWT claims dictate which portal shell is mounted (Student, Teacher, or Admin). The student sidebar "
        "arranges items into Overview, Learning, and Career. The teacher panel renders class walls and grading portals. "
        "The admin control screen organizes approval queues and system ledger logs. While the client determines the route "
        "view dynamically, the backend API enforces access control layers independently."
    )

    add_heading_2("Real-Time Socket Interfacing")
    doc.add_paragraph(
        "During live lectures or proctored assessments, the frontend connects to a real-time gateway using the native WebSockets API. "
        "When the cognitive proctoring agents detect an event (e.g., eye closure threshold exceeded), the orchestrator pushes a socket "
        "packet, triggering immediate warning overlays on the client browser without polling."
    )

    # ----------------------------------------------------
    # SECTION 3
    # ----------------------------------------------------
    add_heading_1("3. Core Backend API Service")
    
    doc.add_paragraph(
        "The core backend is implemented using FastAPI, running asynchronously under a Uvicorn ASGI server. Database transactions "
        "and entities are managed via SQLAlchemy ORM, while Alembic controls schema migration histories."
    )

    add_heading_2("Authentication & Registration Gateway")
    doc.add_paragraph(
        "User registrations hash passwords using bcrypt (12 rounds) via passlib. Login checks query the SQLite/PostgreSQL database, "
        "compare password hashes using constant-time comparison methods, and generate signed JWT tokens containing the user's role "
        "and user ID."
    )

    add_heading_2("Toxicity Moderation & Plagiarism Engine")
    doc.add_paragraph(
        "Class posts undergo asynchronous Content Moderation scans. Homework text files trigger TF-IDF Vectorization "
        "calculations using scikit-learn. The document matrix is compared against all previous student submissions for that task "
        "using Cosine Similarity formulas. This yields a percentage score stored as plagiarism_score, which flags copying infractions."
    )

    add_heading_2("WebRTC Live Proctoring Rooms")
    doc.add_paragraph(
        "FastAPI integrates with the Jitsi Meet IFrame SDK. Teachers can spin up WebRTC video channels. The core API signs JWT room "
        "tokens to restrict room entries to enrolled students."
    )

    # ----------------------------------------------------
    # SECTION 4
    # ----------------------------------------------------
    add_heading_1("4. Cognitive AI Agent Microservices")
    
    doc.add_paragraph(
        "To achieve scalability, each proctoring agent is deployed as a standalone FastAPI microservice. The services communicate "
        "by publishing JSON-serialized event messages to Redis topics."
    )

    # Table for Agents comparison
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Agent Microservice'
    hdr_cells[1].text = 'Primary AI Tools'
    hdr_cells[2].text = 'Key Internal Metric & Action'
    
    for cell in hdr_cells:
        set_cell_bg(cell, "1A365D") # Navy blue header
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        # set text color to white
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
                r.font.size = Pt(10)

    agent_data = [
        ('Attendance Agent', 'MediaPipe Face Mesh, OpenCV', 'Tracks face coordinates. Triggers event if student is absent > N minutes.'),
        ('Alertness Agent', 'MediaPipe Landmarks, NumPy', 'Calculates Eye Aspect Ratio (EAR). Triggers event if EAR < 0.25 for >3 seconds.'),
        ('Content Moderation', 'Detoxify NLP, MobileNet V2', 'Scans posts for toxicity or NSFW attachments. Flags posts if score > 0.70.'),
        ('Screen Integrity', 'YOLOv8, OpenCV', 'Scans screen shares for phones, second persons, or empty chairs during exams.'),
        ('Orchestrator Agent', 'Redis pub/sub, Rules Engine', 'Subscribes to events. Automatically actions events or routes to teacher overrides.')
    ]

    for i, (name, tools, metric) in enumerate(agent_data):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = name
        row_cells[1].text = tools
        row_cells[2].text = metric
        
        # Zebra striping
        bg_color = "F9FAFB" if i % 2 == 0 else "FFFFFF"
        for cell in row_cells:
            set_cell_bg(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            set_cell_borders(cell, top="E2E8F0", bottom="E2E8F0")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)

    doc.add_paragraph("") # Space after table

    # ----------------------------------------------------
    # SECTION 5
    # ----------------------------------------------------
    add_heading_1("5. Event Bus and Relational Database")
    
    doc.add_paragraph(
        "Redis serves as our asynchronous Event Bus. Agents execute `redis.publish('agent_events', payload)` while the Orchestrator "
        "subscribes to this channel. This decouples the core API's uptime from the agents' processing loops."
    )

    add_heading_2("Relational Schema Ledger")
    doc.add_paragraph(
        "The relational database contains 14 tables (Users, Classrooms, Enrollments, WallPosts, Comments, Assignments, Submissions, "
        "Grades, ProctorSessions, ProctorAlerts, AuditLogs, SystemSettings, QuizSources, QuizQuestions)."
    )

    add_heading_2("Security & Immutable Audit Ledger")
    doc.add_paragraph(
        "To ensure compliance, all teacher overrides and agent logs are written to an immutable Audit Log. Each row contains "
        "a SHA-256 hash calculated from its data plus the hash of the preceding row. Manually modifying a row in PostgreSQL breaks the "
        "cryptographic link chain, flagging tampered logs instantly."
    )

    # ----------------------------------------------------
    # SECTION 6
    # ----------------------------------------------------
    add_heading_1("6. Deployment and DevOps")
    
    doc.add_paragraph(
        "All services are containerized using Docker. A central `docker-compose.yml` orchestrates all 7 microservices along with "
        "PostgreSQL and Redis containers. Health checks prevent Dependent services (like FastAPI) from starting up before the "
        "database reports ready. Local volume mounts map folders inside container runtimes to allow uvicorn's `--reload` parameter to "
        "hot-reload code edits dynamically."
    )

    # 7. Document Footer and Page Numbers
    # Add page numbers and section headers to header/footer
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "EngageAI Multi-Agent LMS Portal Specification"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = COLOR_MUTED
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Confidential - For Internal Use Only"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.runs[0].font.size = Pt(8.5)
        fp.runs[0].font.color.rgb = COLOR_MUTED

    doc.save("System_Architecture_Specification.docx")
    print("SUCCESS: System_Architecture_Specification.docx successfully created.")

if __name__ == "__main__":
    create_specification_document()
